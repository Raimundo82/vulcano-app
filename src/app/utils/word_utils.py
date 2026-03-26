import os
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from flask import current_app


def generate_word_with_table(faturas):
    """
    Gera um ficheiro Word (.docx) com uma tabela de faturas.

    Args:
        faturas (list): Lista de dicionários contendo os dados das faturas.

    Returns:
        bytes: Conteúdo do Word gerado.
    """

    document = Document()

    # -------------------------
    # IMAGEM
    # -------------------------
    img_path = os.path.join(current_app.static_folder, "imgs", "SI_faturas.png")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(img_path, width=Inches(1.0))

    document.add_paragraph("")

    # -------------------------
    # TÍTULO
    # -------------------------
    title = document.add_heading("RELATÓRIO RECEPÇÃO DE MATERIAL/SERVIÇO", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    # -------------------------
    # TABELA PROCESSO / FORNECEDOR
    # -------------------------
    info_table = document.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"

    info_table.rows[0].cells[0].text = "N.º PROCESSO DESPESA:"
    info_table.rows[0].cells[1].text = "FORNECEDOR:"
    info_table.rows[1].cells[0].text = "3025006649"
    info_table.rows[1].cells[
        1
    ].text = "252525 - MEO - SERVIÇOS DE COMUNICAÇÕES E MULTIMÉDIA, S.A."

    document.add_paragraph("")

    # -------------------------
    # TEXTO INTRODUTÓRIO
    # -------------------------
    document.add_paragraph(
        "Considera-se o serviço foi realizado de acordo com os termos e condições "
        "contratualizados, tendo sido efetuada uma análise à média dos consumos "
        "por cada conta (UEO), considerando-se que estão em condições para "
        "considerar cumpridos os parâmetros para a receção qualitativa e "
        "quantitativa das seguintes faturas:"
    )

    document.add_paragraph("")

    # -------------------------
    # TABELA DE FATURAS
    # -------------------------
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = [
        "Nº Fatura",
        "Nº Conta",
        "Mês",
        "Ano",
        "Valor em Divida",
        "Valor a Pagar",
    ]
    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    total_amount = 0
    total_amount_to_pay = 0

    for fatura in faturas:
        row = table.add_row().cells
        amount = float(fatura["total_amount"])

        amount_to_pay = fatura["amount_to_pay"]

        try:
            amount_to_pay = float(amount_to_pay or 0)
        except (TypeError, ValueError):
            amount_to_pay = 0.0

        total_amount_to_pay += amount_to_pay
        total_amount += amount

        row[0].text = str(fatura["invoice_number"])
        row[1].text = str(fatura["account_number"])
        row[2].text = str(fatura["invoice_period_month"])
        row[3].text = str(fatura["invoice_period_year"])

        row[4].text = (
            f"€ {amount_to_pay:,.2f}".replace(",", "X")
            .replace(".", ",")
            .replace("X", ".")
        )

        row[5].text = (
            f"€ {amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        )

    # Linha de Total
    total_row = table.add_row().cells
    total_row[3].text = "Total"
    total_row[4].text = (
        f"€ {total_amount_to_pay:,.2f}".replace(",", "X")
        .replace(".", ",")
        .replace("X", ".")
    )
    total_row[5].text = (
        f"€ {total_amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    )

    document.add_paragraph("")
    document.add_paragraph("")

    # -------------------------
    # DATA
    # -------------------------
    data_atual = datetime.now().strftime("%d/%m/%Y")
    document.add_paragraph(f"Superintendência da Informação, {data_atual},")

    document.add_paragraph("")

    # -------------------------
    # ASSINATURAS
    # -------------------------
    footer_table = document.add_table(rows=2, cols=2)
    footer_table.style = "Table Grid"

    footer_table.rows[0].cells[0].text = "O Gestor do Contrato"
    footer_table.rows[0].cells[1].text = "O Chefe de Divisão"

    footer_table.rows[1].cells[0].text = "Ass: ___________________________"
    footer_table.rows[1].cells[1].text = "Ass: ___________________________"

    # -------------------------
    # GUARDAR EM BUFFER
    # -------------------------
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
